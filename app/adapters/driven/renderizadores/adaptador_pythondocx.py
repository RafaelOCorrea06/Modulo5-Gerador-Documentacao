import base64
from io import BytesIO

from docx import Document
from docx.shared import Inches

from app.domain.entidades.artefato import ArtefatoGerado
from app.domain.entidades.documento import DocumentoTecnico


class AdaptadorPythonDocx:
    """
    Adaptador responsável por transformar um DocumentoTecnico em arquivo DOCX.

    Este adaptador usa a biblioteca python-docx para montar o documento.
    Ele recebe a entidade de domínio DocumentoTecnico e devolve um ArtefatoGerado.
    """

    def renderizar(self, documento: DocumentoTecnico) -> ArtefatoGerado:
        doc = Document()

        doc.add_heading(documento.titulo, level=0)

        if documento.subtitulo:
            doc.add_heading(documento.subtitulo, level=1)

        if documento.autor:
            paragrafo_autor = doc.add_paragraph()
            paragrafo_autor.add_run("Autor: ").bold = True
            paragrafo_autor.add_run(documento.autor)

        if documento.metadados:
            doc.add_heading("Metadados", level=1)

            for chave, valor in documento.metadados.items():
                paragrafo = doc.add_paragraph(style="List Bullet")
                paragrafo.add_run(f"{chave}: ").bold = True
                paragrafo.add_run(valor)

        for secao in documento.secoes:
            doc.add_heading(secao.titulo, level=1)

            for paragrafo in secao.paragrafos:
                doc.add_paragraph(paragrafo)

            for lista in secao.listas:
                for item in lista:
                    doc.add_paragraph(item, style="List Bullet")

            for imagem in secao.imagens:
                imagem_bytes = base64.b64decode(imagem.conteudo_base64)
                imagem_stream = BytesIO(imagem_bytes)

                doc.add_picture(imagem_stream, width=Inches(5.5))

                if imagem.legenda:
                    legenda = doc.add_paragraph()
                    legenda.add_run(imagem.legenda).italic = True

        arquivo = BytesIO()
        doc.save(arquivo)
        arquivo.seek(0)

        return ArtefatoGerado(
            nome_arquivo="relatorio-tecnico.docx",
            conteudo=arquivo.getvalue(),
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )
