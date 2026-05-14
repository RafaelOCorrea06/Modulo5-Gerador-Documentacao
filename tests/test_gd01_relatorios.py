from fastapi.testclient import TestClient
from main import app
from io import BytesIO
from docx import Document
import base64
from PIL import Image as PILImage
import time 


client = TestClient(app)

def criar_png_base64_1x1() -> str:
    """
    Cria uma imagem PNG 1x1 em memória e devolve o conteúdo em Base64.

    Isso evita depender de arquivo externo no teste.
    """
    imagem = PILImage.new("RGB", (1, 1), color=(255, 0, 0))

    buffer = BytesIO()
    imagem.save(buffer, format="PNG")

    return base64.b64encode(buffer.getvalue()).decode("utf-8")

def criar_payload_relatorio_grande(formato: str) -> dict:
    """
    Cria um payload grande o suficiente para simular um relatório longo.

    A ideia é gerar muitas seções com parágrafos e listas para aproximar
    um relatório técnico de até 30 páginas, sem depender de arquivo externo.
    """
    secoes = []

    for numero_secao in range(1, 31):
        secoes.append(
            {
                "titulo": f"Seção {numero_secao} - Análise Técnica",
                "paragrafos": [
                    (
                        "Este parágrafo simula conteúdo técnico de um relatório "
                        "de engenharia de software. Ele descreve decisões "
                        "arquiteturais, critérios de implementação, riscos, "
                        "dependências entre serviços e observações relevantes "
                        "para auditoria documental."
                    ),
                    (
                        "Caracteres especiais preservados: ação, código, "
                        "integração, documentação, arquitetura, serviço, "
                        "módulo, validação, rastreabilidade."
                    ),
                    (
                        "O objetivo deste conteúdo é aumentar o volume do "
                        "documento para validar o desempenho da geração em "
                        "relatórios maiores, aproximando o cenário de uso real."
                    ),
                ],
                "listas": [
                    [
                        "Receber estrutura JSON",
                        "Validar seções, parágrafos, listas e imagens",
                        "Gerar artefato no formato solicitado",
                        "Retornar arquivo para download",
                        "Preservar caracteres especiais",
                    ]
                ],
                "imagens": [],
            }
        )

    return {
        "titulo": "Relatório Técnico Grande — GD-01",
        "formato": formato,
        "subtitulo": "Teste de desempenho para relatório extenso",
        "autor": "Rafael",
        "metadados": {
            "Projeto": "Documentação Inteligente",
            "User Story": "GD-01",
            "Critério": "Tempo de geração menor que 10 segundos",
        },
        "secoes": secoes,
    }

def test_root_deve_indicar_servico_rodando():
    response = client.get("/")

    assert response.status_code == 200

    data = response.json()

    assert data["service"] == "gerador-documentacao"
    assert data["status"] == "running"


def test_reports_ping_deve_indicar_modulo_ok():
    response = client.get("/reports/ping")

    assert response.status_code == 200

    data = response.json()

    assert data["module"] == "reports"
    assert data["status"] == "ok"


def test_deve_gerar_relatorio_markdown_com_sucesso():
    payload = {
        "titulo": "Relatório Técnico — GD-01",
        "formato": "md",
        "subtitulo": "Primeira versão em Markdown",
        "autor": "Rafael",
        "metadados": {
            "Projeto": "Documentação Inteligente",
            "User Story": "GD-01",
        },
        "secoes": [
            {
                "titulo": "Objetivo",
                "paragrafos": [
                    "Este relatório testa a geração em Markdown.",
                    "Caracteres especiais: ação, código, integração, lambda, omega, seta.",
                ],
                "listas": [
                    [
                        "Receber JSON",
                        "Gerar Markdown",
                        "Retornar arquivo",
                    ]
                ],
                "imagens": [],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/markdown")
    assert "attachment" in response.headers["content-disposition"]
    assert "relatorio-tecnico.md" in response.headers["content-disposition"]

    conteudo = response.content.decode("utf-8")

    assert "# Relatório Técnico — GD-01" in conteudo
    assert "## Primeira versão em Markdown" in conteudo
    assert "**Autor:** Rafael" in conteudo
    assert "## Metadados" in conteudo
    assert "- **Projeto:** Documentação Inteligente" in conteudo
    assert "## Objetivo" in conteudo
    assert "Caracteres especiais: ação, código, integração, lambda, omega, seta." in conteudo
    assert "- Receber JSON" in conteudo
    assert "- Gerar Markdown" in conteudo

def test_deve_gerar_relatorio_docx_com_sucesso():
    payload = {
        "titulo": "Relatório Técnico — GD-01",
        "formato": "docx",
        "subtitulo": "Primeira versão em DOCX",
        "autor": "Rafael",
        "metadados": {
            "Projeto": "Documentação Inteligente",
            "User Story": "GD-01",
        },
        "secoes": [
            {
                "titulo": "Objetivo",
                "paragrafos": [
                    "Este relatório testa a geração em DOCX.",
                    "Caracteres especiais: ação, código, integração.",
                ],
                "listas": [
                    [
                        "Receber JSON",
                        "Gerar DOCX",
                        "Retornar arquivo",
                    ]
                ],
                "imagens": [],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert "attachment" in response.headers["content-disposition"]
    assert "relatorio-tecnico.docx" in response.headers["content-disposition"]

    assert response.content.startswith(b"PK")

    documento = Document(BytesIO(response.content))
    textos = [paragrafo.text for paragrafo in documento.paragraphs]
    texto_completo = "\n".join(textos)

    assert "Relatório Técnico — GD-01" in texto_completo
    assert "Primeira versão em DOCX" in texto_completo
    assert "Autor: Rafael" in texto_completo
    assert "Projeto: Documentação Inteligente" in texto_completo
    assert "User Story: GD-01" in texto_completo
    assert "Objetivo" in texto_completo
    assert "Este relatório testa a geração em DOCX." in texto_completo
    assert "Caracteres especiais: ação, código, integração." in texto_completo
    assert "Receber JSON" in texto_completo
    assert "Gerar DOCX" in texto_completo  

def test_deve_gerar_relatorio_pdf_com_sucesso():
    payload = {
        "titulo": "Relatório Técnico — GD-01",
        "formato": "pdf",
        "subtitulo": "Primeira versão em PDF",
        "autor": "Rafael",
        "metadados": {
            "Projeto": "Documentação Inteligente",
            "User Story": "GD-01",
        },
        "secoes": [
            {
                "titulo": "Objetivo",
                "paragrafos": [
                    "Este relatório testa a geração em PDF.",
                    "Caracteres especiais: ação, código, integração.",
                ],
                "listas": [
                    [
                        "Receber JSON",
                        "Gerar PDF",
                        "Retornar arquivo",
                    ]
                ],
                "imagens": [],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert "attachment" in response.headers["content-disposition"]
    assert "relatorio-tecnico.pdf" in response.headers["content-disposition"]

    assert response.content.startswith(b"%PDF")
    assert len(response.content) > 1000

def test_deve_gerar_markdown_com_imagem_incorporada():
    imagem_base64 = criar_png_base64_1x1()

    payload = {
        "titulo": "Relatório com Imagem",
        "formato": "md",
        "autor": "Rafael",
        "secoes": [
            {
                "titulo": "Diagrama",
                "paragrafos": [
                    "Esta seção contém uma imagem incorporada em Base64."
                ],
                "listas": [],
                "imagens": [
                    {
                        "nome": "diagrama.png",
                        "conteudo_base64": imagem_base64,
                        "tipo_mime": "image/png",
                        "legenda": "Diagrama simples de arquitetura",
                    }
                ],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/markdown")

    conteudo = response.content.decode("utf-8")

    assert "![Diagrama simples de arquitetura]" in conteudo
    assert "data:image/png;base64," in conteudo
    assert imagem_base64 in conteudo
    assert "*Diagrama simples de arquitetura*" in conteudo


def test_deve_gerar_docx_com_imagem_incorporada():
    imagem_base64 = criar_png_base64_1x1()

    payload = {
        "titulo": "Relatório DOCX com Imagem",
        "formato": "docx",
        "autor": "Rafael",
        "secoes": [
            {
                "titulo": "Diagrama",
                "paragrafos": [
                    "Esta seção contém uma imagem incorporada em Base64."
                ],
                "listas": [],
                "imagens": [
                    {
                        "nome": "diagrama.png",
                        "conteudo_base64": imagem_base64,
                        "tipo_mime": "image/png",
                        "legenda": "Diagrama simples de arquitetura",
                    }
                ],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert response.content.startswith(b"PK")

    documento = Document(BytesIO(response.content))

    textos = [paragrafo.text for paragrafo in documento.paragraphs]
    texto_completo = "\n".join(textos)

    assert "Relatório DOCX com Imagem" in texto_completo
    assert "Diagrama" in texto_completo
    assert "Diagrama simples de arquitetura" in texto_completo
    assert len(documento.inline_shapes) >= 1

def test_deve_gerar_pdf_com_imagem_incorporada():
    imagem_base64 = criar_png_base64_1x1()

    payload = {
        "titulo": "Relatório PDF com Imagem",
        "formato": "pdf",
        "autor": "Rafael",
        "secoes": [
            {
                "titulo": "Diagrama",
                "paragrafos": [
                    "Esta seção contém uma imagem incorporada em Base64."
                ],
                "listas": [],
                "imagens": [
                    {
                        "nome": "diagrama.png",
                        "conteudo_base64": imagem_base64,
                        "tipo_mime": "image/png",
                        "legenda": "Diagrama simples de arquitetura",
                    }
                ],
            }
        ],
    }

    response = client.post("/reports", json=payload)

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.content.startswith(b"%PDF")
    assert len(response.content) > 1000

def test_deve_gerar_relatorio_pdf_grande_em_menos_de_10_segundos():
    payload = criar_payload_relatorio_grande("pdf")

    inicio = time.perf_counter()

    response = client.post("/reports", json=payload)

    duracao = time.perf_counter() - inicio

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.content.startswith(b"%PDF")
    assert len(response.content) > 1000

    assert duracao < 10, (
        f"Geração demorou {duracao:.2f}s, "
        "mas o limite do critério de aceitação é 10s."
    )

def test_deve_gerar_relatorio_docx_grande_em_menos_de_10_segundos():
    payload = criar_payload_relatorio_grande("docx")

    inicio = time.perf_counter()

    response = client.post("/reports", json=payload)

    duracao = time.perf_counter() - inicio

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert response.content.startswith(b"PK")
    assert len(response.content) > 1000

    assert duracao < 10, (
        f"Geração DOCX demorou {duracao:.2f}s, "
        "mas o limite do critério de aceitação é 10s."
    )


def test_deve_gerar_relatorio_markdown_grande_em_menos_de_10_segundos():
    payload = criar_payload_relatorio_grande("md")

    inicio = time.perf_counter()

    response = client.post("/reports", json=payload)

    duracao = time.perf_counter() - inicio

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/markdown")
    assert len(response.content) > 1000

    conteudo = response.content.decode("utf-8")

    assert "# Relatório Técnico Grande — GD-01" in conteudo
    assert "Seção 30 - Análise Técnica" in conteudo

    assert duracao < 10, (
        f"Geração Markdown demorou {duracao:.2f}s, "
        "mas o limite do critério de aceitação é 10s."
    )